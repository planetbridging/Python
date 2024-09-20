import requests
import datetime
import json
import os
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configuration
GITHUB_USERNAME = os.getenv("GITHUB_USER")
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
current_year = datetime.datetime.now().year
CACHE_FILE = "github_commit_cache.json"
EVENTS_PER_PAGE = 100
MAX_PAGES = 3  # Limit the number of pages to fetch to avoid API limits

# Function to load cached commit data
def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, 'r') as f:
            return json.load(f)
    return {}

# Function to save commit data to cache
def save_cache(commit_data):
    with open(CACHE_FILE, 'w') as f:
        json.dump(commit_data, f)

# Function to get commit data from GitHub or cache
def get_commit_data(username):
    # Load cached data
    commit_data = load_cache()

    # Initialize API request parameters
    page = 1
    while page <= MAX_PAGES:
        api_url = f"https://api.github.com/users/{username}/events?page={page}&per_page={EVENTS_PER_PAGE}"
        headers = {
            "Accept": "application/vnd.github.v3+json",
            "Authorization": f"token {GITHUB_TOKEN}"
        }
        
        print(f"Fetching data from GitHub API (Page {page})...")
        response = requests.get(api_url, headers=headers)

        if response.status_code != 200:
            print(f"Failed to fetch data: {response.status_code} - {response.text}")
            break

        events = response.json()

        if not events:
            print(f"No more events found on page {page}.")
            break

        for event in events:
            if event["type"] == "PushEvent":
                repo_name = event["repo"]["name"]
                repo_private = event["repo"].get("private", False)
                visibility = "private" if repo_private else "public"
                
                if repo_name not in commit_data:
                    commit_data[repo_name] = {"commits": [], "total_additions": 0, "total_deletions": 0, "visibility": visibility}
                
                for commit in event["payload"]["commits"]:
                    commit_date = datetime.datetime.strptime(event["created_at"], "%Y-%m-%dT%H:%M:%SZ")
                    if commit_date.year == current_year:
                        commit_url = commit["url"]
                        commit_details = requests.get(commit_url, headers=headers).json()

                        # Extract additions, deletions, and handle missing keys
                        additions = commit_details.get("stats", {}).get("additions", 0)
                        deletions = commit_details.get("stats", {}).get("deletions", 0)
                        html_url = commit_details.get("html_url", "URL not available")

                        commit_data[repo_name]["commits"].append({
                            "sha": commit["sha"],
                            "date": commit_date.strftime("%Y-%m-%d"),
                            "message": commit.get("message", "No commit message"),
                            "additions": additions,
                            "deletions": deletions,
                            "url": html_url
                        })

                        commit_data[repo_name]["total_additions"] += additions
                        commit_data[repo_name]["total_deletions"] += deletions

        page += 1

    # Save updated cache
    save_cache(commit_data)
    return commit_data

# Function to generate daily and monthly graphs
def generate_graphs(commit_data, base_folder):
    overall_dates = []
    overall_additions = []
    overall_deletions = []
    repo_commit_counts = {}

    daily_commit_counts = defaultdict(int)
    monthly_commit_counts = defaultdict(int)

    for repo, data in commit_data.items():
        # Replace directory separators with underscores for safe file names
        safe_repo_name = repo.replace("/", "_")
        repo_folder = os.path.join(base_folder, safe_repo_name)
        os.makedirs(repo_folder, exist_ok=True)

        dates = [commit["date"] for commit in data["commits"]]
        additions = [commit["additions"] for commit in data["commits"]]
        deletions = [commit["deletions"] for commit in data["commits"]]

        # Append to overall data
        overall_dates.extend(dates)
        overall_additions.extend(additions)
        overall_deletions.extend(deletions)

        # Count commits per repository
        repo_commit_counts[repo] = len(data["commits"])

        # Count daily and monthly commits
        for date in dates:
            daily_commit_counts[date] += 1
            month = date[:7]  # YYYY-MM format
            monthly_commit_counts[month] += 1

        # Plot per-repository contributions over time
        plt.figure(figsize=(10, 6))
        plt.plot(dates, additions, label='Additions', color='g')
        plt.plot(dates, deletions, label='Deletions', color='r')
        plt.xticks(rotation=45, ha='right')
        plt.xlabel("Date")
        plt.ylabel("Lines of Code")
        plt.title(f"Lines of Code Committed in {repo} (2024)")
        plt.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(repo_folder, f"{safe_repo_name}_contributions.png"))
        plt.close()

    # Overall activity graph
    plt.figure(figsize=(10, 6))
    plt.plot(overall_dates, overall_additions, label='Additions', color='g')
    plt.plot(overall_dates, overall_deletions, label='Deletions', color='r')
    plt.xticks(rotation=45, ha='right')
    plt.xlabel("Date")
    plt.ylabel("Lines of Code")
    plt.title("Overall Lines of Code Committed in 2024")
    plt.legend()
    plt.tight_layout()
    plt.savefig(os.path.join(base_folder, "overall_contributions.png"))
    plt.close()

    # Commit distribution pie chart
    plt.figure(figsize=(8, 8))
    plt.pie(repo_commit_counts.values(), labels=repo_commit_counts.keys(), autopct='%1.1f%%', startangle=140)
    plt.title("Commit Distribution Across Repositories")
    plt.tight_layout()
    plt.savefig(os.path.join(base_folder, "commit_distribution.png"))
    plt.close()

    # Daily commit activity graph
    sorted_dates = sorted(daily_commit_counts.keys())
    daily_counts = [daily_commit_counts[date] for date in sorted_dates]

    plt.figure(figsize=(10, 6))
    plt.plot(sorted_dates, daily_counts, marker='o', linestyle='-', color='b')
    plt.xticks(rotation=45, ha='right')
    plt.xlabel("Date")
    plt.ylabel("Number of Commits")
    plt.title("Daily Commit Activity in 2024")
    plt.tight_layout()
    plt.savefig(os.path.join(base_folder, "daily_commit_activity.png"))
    plt.close()

    # Monthly commit activity graph
    sorted_months = sorted(monthly_commit_counts.keys())
    monthly_counts = [monthly_commit_counts[month] for month in sorted_months]

    plt.figure(figsize=(10, 6))
    plt.plot(sorted_months, monthly_counts, marker='o', linestyle='-', color='m')
    plt.xticks(rotation=45, ha='right')
    plt.xlabel("Month")
    plt.ylabel("Number of Commits")
    plt.title("Monthly Commit Activity in 2024")
    plt.tight_layout()
    plt.savefig(os.path.join(base_folder, "monthly_commit_activity.png"))
    plt.close()

# Function to create a detailed Word report
def create_word_report(commit_data, base_folder, report_name):
    document = Document()
    document.add_heading(f'{report_name} - GitHub Contributions Report (2024)', 0)

    document.add_heading('Overview', level=1)
    total_commits = sum(len(data["commits"]) for data in commit_data.values())
    total_additions = sum(data["total_additions"] for data in commit_data.values())
    total_deletions = sum(data["total_deletions"] for data in commit_data.values())
    average_additions = total_additions / total_commits if total_commits > 0 else 0
    average_deletions = total_deletions / total_commits if total_commits > 0 else 0

    document.add_paragraph(f'Total Commits: {total_commits}')
    document.add_paragraph(f'Total Lines Added: {total_additions}')
    document.add_paragraph(f'Total Lines Deleted: {total_deletions}')
    document.add_paragraph(f'Average Lines Added per Commit: {average_additions:.2f}')
    document.add_paragraph(f'Average Lines Deleted per Commit: {average_deletions:.2f}')

    # Add overall activity graph
    document.add_heading('Overall Activity', level=2)
    document.add_paragraph('This graph shows the total lines of code added and deleted over time across all repositories.')
    document.add_picture(os.path.join(base_folder, 'overall_contributions.png'), width=Inches(6))

    # Add commit distribution pie chart
    document.add_heading('Commit Distribution', level=2)
    document.add_paragraph('This pie chart shows the distribution of commits across different repositories.')
    document.add_picture(os.path.join(base_folder, 'commit_distribution.png'), width=Inches(6))

    # Add daily commit activity graph
    document.add_heading('Daily Commit Activity', level=2)
    document.add_paragraph('This graph shows the number of commits made each day across all repositories.')
    document.add_picture(os.path.join(base_folder, 'daily_commit_activity.png'), width=Inches(6))

    # Add monthly commit activity graph
    document.add_heading('Monthly Commit Activity', level=2)
    document.add_paragraph('This graph shows the number of commits made each month across all repositories.')
    document.add_picture(os.path.join(base_folder, 'monthly_commit_activity.png'), width=Inches(6))

    for repo, data in commit_data.items():
        safe_repo_name = repo.replace("/", "_")
        repo_folder = os.path.join(base_folder, safe_repo_name)

        document.add_heading(repo, level=2)
        document.add_paragraph(f"Total Commits: {len(data['commits'])}")
        document.add_paragraph(f"Total Lines Added: {data['total_additions']}")
        document.add_paragraph(f"Total Lines Deleted: {data['total_deletions']}")

        # Key statistics
        if data['commits']:
            largest_commit = max(data['commits'], key=lambda c: c['additions'] + c['deletions'])
            most_active_day = max(set([commit['date'] for commit in data['commits']]), key=lambda d: sum(c['date'] == d for c in data['commits']))

            document.add_paragraph(f"Largest Commit: {largest_commit['sha']} on {largest_commit['date']} ({largest_commit['additions']} additions, {largest_commit['deletions']} deletions)")
            document.add_paragraph(f"Most Active Day: {most_active_day}")

        document.add_heading('Commits', level=3)
        for commit in data["commits"]:
            document.add_paragraph(f"Date: {commit['date']}")
            document.add_paragraph(f"Commit: {commit['sha']}")
            document.add_paragraph(f"Message: {commit['message']}")
            document.add_paragraph(f"Additions: {commit['additions']}, Deletions: {commit['deletions']}")
            document.add_paragraph(f"URL: {commit['url']}")
            document.add_paragraph()

        document.add_picture(os.path.join(repo_folder, f"{safe_repo_name}_contributions.png"), width=Inches(6))

    document.save(os.path.join(base_folder, f'{report_name}_GitHub_Contribution_Report.docx'))

# Function to create a summary Word report with only graphs and titles
def create_graphs_only_report(commit_data, base_folder, report_name):
    document = Document()
    document.add_heading(f'{report_name} - GitHub Contributions Graphs (2024)', 0)

    # Add overall activity graph
    document.add_heading('Overall Activity', level=2)
    document.add_picture(os.path.join(base_folder, 'overall_contributions.png'), width=Inches(6))

    # Add commit distribution pie chart
    document.add_heading('Commit Distribution', level=2)
    document.add_picture(os.path.join(base_folder, 'commit_distribution.png'), width=Inches(6))

    # Add daily commit activity graph
    document.add_heading('Daily Commit Activity', level=2)
    document.add_picture(os.path.join(base_folder, 'daily_commit_activity.png'), width=Inches(6))

    # Add monthly commit activity graph
    document.add_heading('Monthly Commit Activity', level=2)
    document.add_picture(os.path.join(base_folder, 'monthly_commit_activity.png'), width=Inches(6))

    for repo, data in commit_data.items():
        safe_repo_name = repo.replace("/", "_")
        repo_folder = os.path.join(base_folder, safe_repo_name)

        document.add_heading(repo, level=2)
        document.add_picture(os.path.join(repo_folder, f"{safe_repo_name}_contributions.png"), width=Inches(6))

    document.save(os.path.join(base_folder, f'{report_name}_GitHub_Graphs_Only_Report.docx'))

# Function to create individual reports for each repository
def create_individual_repo_reports(commit_data, base_folder):
    for repo, data in commit_data.items():
        safe_repo_name = repo.replace("/", "_")
        repo_folder = os.path.join(base_folder, safe_repo_name)
        os.makedirs(repo_folder, exist_ok=True)

        document = Document()
        document.add_heading(f'GitHub Contributions Report - {GITHUB_USERNAME} ({repo}) (2024)', 0)

        document.add_paragraph(f"Total Commits: {len(data['commits'])}")
        document.add_paragraph(f"Total Lines Added: {data['total_additions']}")
        document.add_paragraph(f"Total Lines Deleted: {data['total_deletions']}")

        # Key statistics
        if data['commits']:
            largest_commit = max(data['commits'], key=lambda c: c['additions'] + c['deletions'])
            most_active_day = max(set([commit['date'] for commit in data['commits']]), key=lambda d: sum(c['date'] == d for c in data['commits']))

            document.add_paragraph(f"Largest Commit: {largest_commit['sha']} on {largest_commit['date']} ({largest_commit['additions']} additions, {largest_commit['deletions']} deletions)")
            document.add_paragraph(f"Most Active Day: {most_active_day}")

        document.add_heading('Commits', level=3)
        for commit in data["commits"]:
            document.add_paragraph(f"Date: {commit['date']}")
            document.add_paragraph(f"Commit: {commit['sha']}")
            document.add_paragraph(f"Message: {commit['message']}")
            document.add_paragraph(f"Additions: {commit['additions']}, Deletions: {commit['deletions']}")
            document.add_paragraph(f"URL: {commit['url']}")
            document.add_paragraph()

        document.add_picture(os.path.join(repo_folder, f"{safe_repo_name}_contributions.png"), width=Inches(6))

        document.save(os.path.join(repo_folder, f'{safe_repo_name}_GitHub_Contribution_Report.docx'))

# Execute the functions
commit_data = get_commit_data(GITHUB_USERNAME)

# Separate data for public and private repositories
public_commit_data = {repo: data for repo, data in commit_data.items() if data["visibility"] == "public"}
private_commit_data = {repo: data for repo, data in commit_data.items() if data["visibility"] == "private"}

# Generate graphs and reports for public repositories
os.makedirs("public", exist_ok=True)
generate_graphs(public_commit_data, "public")
create_word_report(public_commit_data, "public", "Public Repositories")
create_graphs_only_report(public_commit_data, "public", "Public Repositories")
create_individual_repo_reports(public_commit_data, "public")

# Generate graphs and reports for private repositories
os.makedirs("private", exist_ok=True)
generate_graphs(private_commit_data, "private")
create_word_report(private_commit_data, "private", "Private Repositories")
create_graphs_only_report(private_commit_data, "private", "Private Repositories")
create_individual_repo_reports(private_commit_data, "private")

# Generate a combined report for all repositories
os.makedirs("combined", exist_ok=True)
generate_graphs(commit_data, "combined")
create_word_report(commit_data, "combined", "All Repositories")
create_graphs_only_report(commit_data, "combined", "All Repositories")
